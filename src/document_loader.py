from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from src.models import SourceDocument


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


def load_documents(raw_dir: Path) -> list[SourceDocument]:
    documents: list[SourceDocument] = []
    for path in sorted(raw_dir.rglob("*")):
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        if path.name.lower() == "readme.md":
            continue
        documents.extend(_load_single_document(path))
    return documents


def _load_single_document(path: Path) -> list[SourceDocument]:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return [_build_document(path, text)]
    if suffix == ".docx":
        doc = DocxDocument(path)
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        return [_build_document(path, text)]
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        results: list[SourceDocument] = []
        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                results.append(_build_document(path, text, page=index))
        return results
    return []


def _build_document(path: Path, text: str, page: int | None = None) -> SourceDocument:
    return SourceDocument(
        source_path=str(path.resolve()),
        source_name=path.name,
        text=text,
        page=page,
    )
